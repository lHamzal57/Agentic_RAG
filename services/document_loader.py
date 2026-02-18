# file_validation.py
import pdfplumber
from docx import Document
import csv
from io import BytesIO, StringIO
import mimetypes
import pdfplumber
from docx import Document

def extract_text(file_bytes: bytes, mime_type: str) -> str:
    if mime_type == "application/pdf":
        text = []
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text.append(page.extract_text() or "")
        return "\n".join(text)

    if mime_type.endswith("wordprocessingml.document"):
        doc = Document(BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs)

    if mime_type == "text/csv":
        csv_text = StringIO(file_bytes.decode("utf-8"))
        reader = csv.reader(csv_text)
        return "\n".join(",".join(row) for row in reader)

    return ""

def load_document_text(path: str) -> str:
    mime_type, _ = mimetypes.guess_type(path)

    if not mime_type:
        raise ValueError("Unable to detect file type")

    # ---------- PDF ----------
    if mime_type == "application/pdf":
        return extract_pdf_text(path)

    # ---------- DOCX ----------
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_docx_text(path)

    # ---------- Plain text / markdown ----------
    if mime_type.startswith("text/"):
        return extract_text_file(path)

    raise ValueError(f"Unsupported file type: {mime_type}")

def extract_docx_text(path):
    doc = Document(path)
    return "\n".join(
        p.text for p in doc.paragraphs if p.text.strip()
    )

def extract_pdf_text(path):
    text = ""
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                text += f"\n[Page {i+1}]\n{page_text}"
    return text

def extract_text_file(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()



